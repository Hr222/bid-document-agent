from __future__ import annotations

import base64
from pathlib import Path

import pytest

from app.interfaces.http.schemas import ParsedBlock, ParsedDocumentResult
from app.infrastructure.ocr.tencent_ocr import PolicyOcrService
from app.modules.ingestion.pipeline.steps.policy_parser import PolicyParserService
from app.modules.ingestion.pipeline.steps.policy_text_assembler import PolicyTextAssemblerService


def make_block(
    *,
    block_id: str,
    order: int,
    block_type: str,
    text: str | None = None,
    page_no: int | None = None,
    metadata: dict | None = None,
) -> ParsedBlock:
    return ParsedBlock(
        block_id=block_id,
        order=order,
        page_no=page_no,
        block_type=block_type,
        source="direct",
        text=text,
        metadata=metadata or {},
        layout_hint={},
    )


def test_docx_with_text_and_inline_image_still_selects_image_for_ocr() -> None:
    document = ParsedDocumentResult(
        parser_status="parsed",
        source_path="D:/workspace/bid-document-agent/tests/fixtures/mixed.docx",
        file_type="docx",
        suspected_scanned=False,
        blocks=[
            make_block(
                block_id="text-1",
                order=0,
                block_type="text",
                text="AAA[IMAGE_OCR_1]BBB",
                metadata={
                    "container": "paragraph",
                    "ocr_placeholder_tokens": ["[IMAGE_OCR_1]"],
                    "has_effective_text": True,
                },
            ),
            make_block(
                block_id="image-1",
                order=1,
                block_type="image",
                metadata={
                    "container": "paragraph",
                    "placeholder_token": "[IMAGE_OCR_1]",
                    "image_bytes": "89504e47",
                    "image_name": "inline.png",
                },
            ),
        ],
        notes=[],
    )

    targets = PolicyOcrService(client=object())._select_ocr_targets(document)

    assert [block.block_id for block in targets] == ["image-1"]


def test_mixed_pdf_with_embedded_image_selects_image_for_ocr() -> None:
    document = ParsedDocumentResult(
        parser_status="parsed",
        source_path="D:/workspace/bid-document-agent/tests/fixtures/mixed.pdf",
        file_type="pdf",
        suspected_scanned=False,
        blocks=[
            make_block(block_id="text-1", order=0, block_type="text", text="visible text"),
            make_block(
                block_id="image-1",
                order=1,
                block_type="image",
                page_no=1,
                metadata={"image_bytes": "89504e47", "image_name": "stamp.png"},
            ),
        ],
        notes=[],
    )

    targets = PolicyOcrService(client=object())._select_ocr_targets(document)

    assert [block.block_id for block in targets] == ["image-1"]


def test_parse_docx_keeps_inline_image_placeholder_order(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document

    image_path = tmp_path / "inline.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wn2bWAAAAAASUVORK5CYII="
        )
    )

    source_path = tmp_path / "inline.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("AAA")
    paragraph.add_run().add_picture(str(image_path))
    paragraph.add_run("BBB")
    document.save(source_path)

    result = PolicyParserService().parse_document(
        source_path=str(source_path),
        parse_method="direct",
    )

    text_blocks = [block for block in result.blocks if block.block_type == "text"]
    image_blocks = [block for block in result.blocks if block.block_type == "image"]

    assert len(text_blocks) == 1
    assert len(image_blocks) == 1
    assert image_blocks[0].metadata["placeholder_token"] in text_blocks[0].text
    assert text_blocks[0].text == f"AAA{image_blocks[0].metadata['placeholder_token']}BBB"


def test_text_assembler_reinserts_inline_ocr_text_and_hides_placeholder() -> None:
    document = ParsedDocumentResult(
        parser_status="parsed",
        source_path="D:/workspace/bid-document-agent/tests/fixtures/mixed.docx",
        file_type="docx",
        suspected_scanned=False,
        blocks=[
            make_block(
                block_id="text-1",
                order=0,
                block_type="text",
                text="AAA[IMAGE_OCR_1]BBB",
                metadata={
                    "container": "paragraph",
                    "ocr_placeholder_tokens": ["[IMAGE_OCR_1]"],
                    "has_effective_text": True,
                },
            ),
            make_block(
                block_id="image-1",
                order=1,
                block_type="image",
                text="插图识别文本",
                metadata={
                    "container": "paragraph",
                    "placeholder_token": "[IMAGE_OCR_1]",
                    "image_name": "inline.png",
                },
            ),
        ],
        notes=[],
    )

    result = PolicyTextAssemblerService().assemble(document, parse_method="mixed")

    assert result.raw_text == "AAA插图识别文本BBB"
    assert result.paragraphs == ["AAA插图识别文本BBB"]
    assert all("[IMAGE_OCR_" not in line.text for line in result.lines)
