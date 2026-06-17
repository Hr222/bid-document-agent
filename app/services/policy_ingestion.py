from __future__ import annotations

from collections import Counter
from datetime import UTC, datetime
from hashlib import sha256
from pathlib import Path
import re

from app.schemas.policy_ingestion import (
    PolicyCandidateItem,
    PolicyPreviewResponse,
    PolicyScanRequest,
    PolicyScanResponse,
    PolicyScanStats,
)


class PolicyIngestionService:
    _excluded_extensions = {".jpg", ".jpeg", ".png", ".zip", ".rar", ".7z"}
    _direct_parse_extensions = {".docx", ".pdf", ".txt"}
    _review_extensions = {".doc"}
    _excluded_keywords = {
        "身份证",
        "签字",
        "签名",
        "公章",
        "法人章",
        "印章",
        "账号",
        "密码",
        "联系方式",
        "联系",
        "logo",
        "照片",
        "截图",
        "盖章",
        "空白",
        "模板",
    }
    _scanned_keywords = {"扫描", "图片版", "盖章", "签字", "签名"}

    def scan_candidates(self, request: PolicyScanRequest) -> PolicyScanResponse:
        root = Path(request.source_root)
        if not root.exists():
            raise FileNotFoundError(f"Source root does not exist: {root}")
        if not root.is_dir():
            raise NotADirectoryError(f"Source root is not a directory: {root}")

        files = sorted(path for path in root.rglob("*") if path.is_file())
        counts = Counter(path.suffix.lower() or "<none>" for path in files)
        candidates = [self._build_candidate(root=root, path=path) for path in files[: request.limit]]

        stats = PolicyScanStats(
            total_files=len(files),
            included_files=sum(1 for item in candidates if item.recommended_action == "include"),
            excluded_files=sum(1 for item in candidates if item.recommended_action == "exclude"),
            review_files=sum(1 for item in candidates if item.recommended_action == "review"),
            by_extension=dict(sorted(counts.items())),
        )
        return PolicyScanResponse(
            source_root=str(root),
            scanned_at=datetime.now(UTC),
            stats=stats,
            candidates=candidates,
        )

    def preview_parse(self, source_path: str) -> PolicyPreviewResponse:
        path = Path(source_path)
        if not path.exists():
            raise FileNotFoundError(f"Source file does not exist: {path}")
        if not path.is_file():
            raise IsADirectoryError(f"Source path is not a file: {path}")

        candidate = self._build_candidate(root=path.parent, path=path)
        notes: list[str] = []
        raw_text = ""
        clean_text = ""
        page_count: int | None = None
        parser_status = "pending"

        if candidate.recommended_action == "exclude":
            notes.append(candidate.exclude_reason or "File excluded by hard filter.")
        elif candidate.parse_method == "manual":
            parser_status = "failed"
            notes.append("Legacy .doc is not parsed automatically yet. Convert it to .docx first.")
        elif candidate.parse_method == "ocr":
            notes.append("This file looks like a scanned or stamped copy and likely needs OCR.")

        if candidate.parse_method == "direct":
            parser_status, raw_text, page_count, parse_notes = self._extract_text(path)
            notes.extend(parse_notes)
            clean_text = self._clean_text(raw_text)
            if path.suffix.lower() == ".pdf" and len(clean_text.strip()) < 50:
                candidate.suspected_scanned = True
                notes.append("PDF extracted very little text; this may be an image-only scan.")
            if not clean_text.strip():
                notes.append("No usable text extracted.")

        text_preview = clean_text[:1200]
        detected_titles = self._detect_titles(clean_text)
        return PolicyPreviewResponse(
            source_path=str(path),
            file_name=path.name,
            extension=path.suffix.lower(),
            parse_method=candidate.parse_method,
            parser_status=parser_status,
            suspected_scanned=candidate.suspected_scanned,
            page_count=page_count,
            raw_text_chars=len(raw_text),
            clean_text_chars=len(clean_text),
            text_preview=text_preview,
            detected_titles=detected_titles,
            notes=notes,
        )

    def _build_candidate(self, root: Path, path: Path) -> PolicyCandidateItem:
        extension = path.suffix.lower()
        relative_path = str(path.relative_to(root)) if path != root else path.name
        full_text = f"{path.name} {relative_path}"
        suspected_scanned = self._contains_keyword(full_text, self._scanned_keywords)

        recommended_action = "review"
        parse_method = "manual"
        include_reason: str | None = None
        exclude_reason: str | None = None

        if path.stat().st_size == 0:
            recommended_action = "exclude"
            parse_method = "skip"
            exclude_reason = "Empty file."
        elif extension in self._excluded_extensions:
            recommended_action = "exclude"
            parse_method = "skip"
            exclude_reason = "Image or archive file is excluded in the first batch."
        elif self._contains_keyword(full_text, self._excluded_keywords):
            recommended_action = "exclude"
            parse_method = "skip"
            exclude_reason = "Matched the first-batch exclusion keyword rule."
        elif extension in self._review_extensions:
            recommended_action = "review"
            parse_method = "manual"
            include_reason = "Legacy Word file needs conversion or a dedicated parser."
        elif extension in self._direct_parse_extensions:
            if suspected_scanned:
                recommended_action = "review"
                parse_method = "ocr" if extension == ".pdf" else "direct"
                include_reason = "Parsable file, but it looks like a scanned or stamped copy."
            else:
                recommended_action = "include"
                parse_method = "direct"
                include_reason = "Native text file suitable for first-batch ingestion."
        else:
            recommended_action = "exclude"
            parse_method = "skip"
            exclude_reason = "Unsupported file type for the first batch."

        return PolicyCandidateItem(
            source_path=str(path),
            relative_path=relative_path,
            file_name=path.name,
            extension=extension,
            size_bytes=path.stat().st_size,
            sha256=self._file_hash(path),
            recommended_action=recommended_action,
            parse_method=parse_method,
            suspected_scanned=suspected_scanned,
            policy_name_guess=self._guess_policy_name(path.stem),
            include_reason=include_reason,
            exclude_reason=exclude_reason,
        )

    def _extract_text(self, path: Path) -> tuple[str, str, int | None, list[str]]:
        notes: list[str] = []
        extension = path.suffix.lower()

        if extension == ".docx":
            try:
                from docx import Document
            except ImportError:
                notes.append("Missing dependency: python-docx is not installed.")
                return "failed", "", None, notes

            document = Document(path)
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_rows: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        table_rows.append(" | ".join(values))
            text = "\n".join(paragraphs + table_rows)
            return "parsed", text, None, notes

        if extension == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                notes.append("Missing dependency: pypdf is not installed.")
                return "failed", "", None, notes

            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages:
                pages.append(page.extract_text() or "")
            return "parsed", "\n".join(pages), len(reader.pages), notes

        if extension == ".txt":
            for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
                try:
                    return "parsed", path.read_text(encoding=encoding), None, notes
                except UnicodeDecodeError:
                    continue
            notes.append("Unable to decode txt with utf-8 or gbk-compatible encodings.")
            return "failed", "", None, notes

        notes.append("No parser configured for this file type yet.")
        return "failed", "", None, notes

    def _clean_text(self, raw_text: str) -> str:
        normalized = raw_text.replace("\r\n", "\n").replace("\r", "\n").replace("\u3000", " ")
        normalized = re.sub(r"[ \t]+", " ", normalized)
        lines = [line.strip() for line in normalized.split("\n")]

        cleaned_lines: list[str] = []
        blank_streak = 0
        for line in lines:
            if not line:
                blank_streak += 1
                if blank_streak <= 1:
                    cleaned_lines.append("")
                continue

            blank_streak = 0
            if re.fullmatch(r"第?\s*\d+\s*页", line) or re.fullmatch(r"-\s*\d+\s*-", line):
                continue
            cleaned_lines.append(line)

        return "\n".join(cleaned_lines).strip()

    def _detect_titles(self, clean_text: str) -> list[str]:
        if not clean_text:
            return []

        title_patterns = (
            r"^第[一二三四五六七八九十百千0-9]+章.*$",
            r"^第[一二三四五六七八九十百千0-9]+节.*$",
            r"^第[一二三四五六七八九十百千0-9]+条.*$",
            r"^[一二三四五六七八九十]+、.*$",
            r"^\([0-9]+\).*$",
        )
        compiled = [re.compile(pattern) for pattern in title_patterns]

        titles: list[str] = []
        for line in clean_text.splitlines():
            if any(pattern.match(line) for pattern in compiled):
                titles.append(line)
            if len(titles) >= 20:
                break
        return titles

    def _guess_policy_name(self, stem: str) -> str:
        cleaned = re.sub(r"^\d{6,8}", "", stem)
        cleaned = re.sub(r"[（(][^）)]*(审核|盖章|空白|模板)[）)]", "", cleaned)
        cleaned = re.sub(r"[-_]{2,}", "-", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned)
        return cleaned.strip(" -_") or stem

    def _contains_keyword(self, text: str, keywords: set[str]) -> bool:
        lowered = text.lower()
        return any(keyword.lower() in lowered for keyword in keywords)

    def _file_hash(self, path: Path) -> str:
        digest = sha256()
        with path.open("rb") as file:
            for chunk in iter(lambda: file.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()
