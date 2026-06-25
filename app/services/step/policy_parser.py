from __future__ import annotations

import re
from pathlib import Path

from app.schemas.policy_pipeline import ParsedTextResult, ParseRoutingResult

CHINESE_HEADING_PATTERNS = (
    r"^第[一二三四五六七八九十百千万零〇两0-9]+章.*$",
    r"^第[一二三四五六七八九十百千万零〇两0-9]+节.*$",
    r"^第[一二三四五六七八九十百千万零〇两0-9]+条.*$",
)


class PolicyParserService:
    """
    负责制度文件的解析阶段。

    - 步骤 4：解析器选择
    - 步骤 5：原文抽取
    """

    def route_parser(self, source_path: str) -> ParseRoutingResult:
        """
        决定当前文件应交给哪个解析器处理。

        这是流水线的第 4 步，应该保持轻量；真正的文本抽取发生在 `parse`。
        """
        path = Path(source_path)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return ParseRoutingResult(
                parser_name="DocxParser",
                parse_method="docx",
                suspected_scanned_pdf=False,
                notes=[],
            )
        if suffix == ".pdf":
            return ParseRoutingResult(
                parser_name="PdfParser",
                parse_method="pdf",
                suspected_scanned_pdf=False,
                notes=[],
            )

        raise ValueError(f"当前文件类型未配置解析器：{suffix}")

    def parse(self, *, source_path: str, parse_method: str) -> ParsedTextResult:
        """步骤 5：抽取原始文本，并附带解析阶段识别出的结构提示。"""
        if parse_method == "docx":
            return self._parse_docx(source_path)
        if parse_method == "pdf":
            return self._parse_pdf(source_path)
        raise ValueError(f"不支持的解析方式：{parse_method}")

    def _parse_docx(self, source_path: str) -> ParsedTextResult:
        """步骤 5.1：将 `.docx` 文件解析成段落和表格行文本。"""
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("缺少依赖：未安装 python-docx。") from exc

        document = Document(source_path)
        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        tables: list[str] = []
        for table in document.tables:
            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_values:
                    tables.append(" | ".join(row_values))

        raw_text = "\n".join(paragraphs + tables)
        return ParsedTextResult(
            parser_status="parsed",
            source_path=source_path,
            raw_text=raw_text,
            page_count=None,
            suspected_scanned=False,
            paragraphs=paragraphs,
            tables=tables,
            title_candidates=self._detect_title_candidates(paragraphs),
            notes=[],
        )

    def _parse_pdf(self, source_path: str) -> ParsedTextResult:
        """步骤 5.2：解析可抽文本 PDF，并标记疑似图片扫描件。"""
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("缺少依赖：未安装 pypdf。") from exc

        reader = PdfReader(source_path)
        pages: list[str] = []
        notes: list[str] = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")

        raw_text = "\n".join(pages)
        paragraphs = [line.strip() for line in raw_text.splitlines() if line.strip()]
        non_whitespace_chars = len(re.sub(r"\s+", "", raw_text))
        suspected_scanned = non_whitespace_chars < 80
        if suspected_scanned:
            notes.append(
                "提取文本过短，当前 MVP 将该 PDF 视为疑似扫描件。"
            )

        return ParsedTextResult(
            parser_status="parsed",
            source_path=source_path,
            raw_text=raw_text,
            page_count=len(reader.pages),
            suspected_scanned=suspected_scanned,
            paragraphs=paragraphs,
            tables=[],
            title_candidates=self._detect_title_candidates(paragraphs),
            notes=notes,
        )

    def _detect_title_candidates(self, paragraphs: list[str]) -> list[str]:
        """步骤 5 的辅助动作：收集疑似标题行，辅助后续章节拆分。"""
        compiled = [re.compile(pattern) for pattern in CHINESE_HEADING_PATTERNS]

        titles: list[str] = []
        for paragraph in paragraphs:
            if any(pattern.match(paragraph) for pattern in compiled):
                titles.append(paragraph)
            if len(titles) >= 30:
                break
        return titles
