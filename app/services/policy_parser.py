from __future__ import annotations

from pathlib import Path
import re

from app.schemas.policy_pipeline import ParsedTextResult, ParseRoutingResult


class PolicyParserService:
    """Choose a parser and extract raw content from normalized files."""

    def route_parser(self, source_path: str) -> ParseRoutingResult:
        """
        Decide which parser should handle the current file.

        This is step 4 of the pipeline and should stay lightweight. The actual
        extraction happens in `parse`.
        """
        path = Path(source_path)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return ParseRoutingResult(
                parser_name="DocxParser",
                parse_method="docx",
                suspected_scanned_pdf=False,
                notes=[],
            )
        if suffix == ".pdf":
            return ParseRoutingResult(
                parser_name="PdfParser",
                parse_method="pdf",
                suspected_scanned_pdf=False,
                notes=[],
            )

        raise ValueError(f"No parser configured for file type: {suffix}")

    def parse(self, *, source_path: str, parse_method: str) -> ParsedTextResult:
        """Extract raw text plus parser-side structure hints."""
        if parse_method == "docx":
            return self._parse_docx(source_path)
        if parse_method == "pdf":
            return self._parse_pdf(source_path)
        raise ValueError(f"Unsupported parse method: {parse_method}")

    def _parse_docx(self, source_path: str) -> ParsedTextResult:
        """Parse a `.docx` file into paragraphs and table rows."""
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Missing dependency: python-docx is not installed.") from exc

        document = Document(source_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        tables: list[str] = []
        for table in document.tables:
            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_values:
                    tables.append(" | ".join(row_values))

        raw_text = "\n".join(paragraphs + tables)
        return ParsedTextResult(
            parser_status="success",
            source_path=source_path,
            raw_text=raw_text,
            page_count=None,
            paragraphs=paragraphs,
            tables=tables,
            title_candidates=self._detect_title_candidates(paragraphs),
            notes=[],
        )

    def _parse_pdf(self, source_path: str) -> ParsedTextResult:
        """Parse a text PDF and flag potential image-only scans."""
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("Missing dependency: pypdf is not installed.") from exc

        reader = PdfReader(source_path)
        pages: list[str] = []
        notes: list[str] = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")

        raw_text = "\n".join(pages)
        paragraphs = [line.strip() for line in raw_text.splitlines() if line.strip()]
        if len("".join(paragraphs)) < 80:
            notes.append("Extracted text is very short; this PDF may be scan-based and need OCR later.")

        return ParsedTextResult(
            parser_status="success",
            source_path=source_path,
            raw_text=raw_text,
            page_count=len(reader.pages),
            paragraphs=paragraphs,
            tables=[],
            title_candidates=self._detect_title_candidates(paragraphs),
            notes=notes,
        )

    def _detect_title_candidates(self, paragraphs: list[str]) -> list[str]:
        """Collect section-like lines to help downstream structure splitting."""
        patterns = (
            r"^第[一二三四五六七八九十百千0-9]+章.*$",
            r"^第[一二三四五六七八九十百千0-9]+节.*$",
            r"^第[一二三四五六七八九十百千0-9]+条.*$",
            r"^[一二三四五六七八九十]+、.*$",
        )
        compiled = [re.compile(pattern) for pattern in patterns]

        titles: list[str] = []
        for paragraph in paragraphs:
            if any(pattern.match(paragraph) for pattern in compiled):
                titles.append(paragraph)
            if len(titles) >= 30:
                break
        return titles
